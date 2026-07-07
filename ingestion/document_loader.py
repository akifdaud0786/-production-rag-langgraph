"""
Document loading layer.

Supports local files (PDF, DOCX, HTML, TXT, MD) and, when GCP credentials are
configured, reading directly from a GCS "raw" bucket and writing normalized
text back to a "processed" bucket — mirroring the Ingestion Pipeline (GCP)
box in the reference architecture.
"""
from __future__ import annotations

import logging
import os
from dataclasses import dataclass, field
from pathlib import Path
from typing import Iterable, Iterator

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".html", ".htm", ".txt", ".md"}


@dataclass
class RawDocument:
    """A single loaded document prior to chunking."""
    doc_id: str
    source_path: str
    text: str
    metadata: dict = field(default_factory=dict)


def _load_pdf(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    return "\n\n".join(page.extract_text() or "" for page in reader.pages)


def _load_docx(path: Path) -> str:
    import docx

    document = docx.Document(str(path))
    return "\n".join(p.text for p in document.paragraphs)


def _load_html(path: Path) -> str:
    from bs4 import BeautifulSoup

    soup = BeautifulSoup(path.read_text(encoding="utf-8", errors="ignore"), "html.parser")
    for tag in soup(["script", "style", "nav", "footer"]):
        tag.decompose()
    return soup.get_text(separator="\n")


def _load_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


_LOADERS = {
    ".pdf": _load_pdf,
    ".docx": _load_docx,
    ".html": _load_html,
    ".htm": _load_html,
    ".txt": _load_text,
    ".md": _load_text,
}


def load_local_documents(source_dir: str | Path) -> Iterator[RawDocument]:
    """Yield RawDocument objects for every supported file under source_dir."""
    source_dir = Path(source_dir)
    if not source_dir.exists():
        raise FileNotFoundError(f"Source directory not found: {source_dir}")

    for path in sorted(source_dir.rglob("*")):
        if path.suffix.lower() not in SUPPORTED_EXTENSIONS or not path.is_file():
            continue
        loader = _LOADERS[path.suffix.lower()]
        try:
            text = loader(path)
        except Exception as exc:  # noqa: BLE001
            logger.warning("Failed to load %s: %s", path, exc)
            continue
        if not text.strip():
            continue
        yield RawDocument(
            doc_id=path.stem,
            source_path=str(path),
            text=text,
            metadata={"filename": path.name, "extension": path.suffix.lower()},
        )


def load_from_gcs(bucket_name: str, prefix: str = "") -> Iterator[RawDocument]:
    """Stream documents directly from a GCS raw bucket (production ingestion path)."""
    from google.cloud import storage

    client = storage.Client()
    bucket = client.bucket(bucket_name)
    for blob in bucket.list_blobs(prefix=prefix):
        ext = Path(blob.name).suffix.lower()
        if ext not in SUPPORTED_EXTENSIONS:
            continue
        local_tmp = Path(f"/tmp/{Path(blob.name).name}")
        blob.download_to_filename(str(local_tmp))
        loader = _LOADERS[ext]
        try:
            text = loader(local_tmp)
        finally:
            local_tmp.unlink(missing_ok=True)
        if not text.strip():
            continue
        yield RawDocument(
            doc_id=Path(blob.name).stem,
            source_path=f"gs://{bucket_name}/{blob.name}",
            text=text,
            metadata={"filename": blob.name, "extension": ext, "bucket": bucket_name},
        )


def write_processed_to_gcs(bucket_name: str, doc: RawDocument) -> str:
    """Write normalized/cleaned text back to the processed GCS bucket."""
    from google.cloud import storage

    client = storage.Client()
    bucket = client.bucket(bucket_name)
    blob_name = f"{doc.doc_id}.txt"
    blob = bucket.blob(blob_name)
    blob.upload_from_string(doc.text, content_type="text/plain")
    return f"gs://{bucket_name}/{blob_name}"
